#!/usr/bin/env python3

import argparse
import re
import sys

from azure.identity import DefaultAzureCredential
from azure.mgmt.loganalytics import LogAnalyticsManagementClient
from azure.mgmt.securityinsight import SecurityInsights
from docx import Document
from docx.shared import Inches


def get_query_tables(kql_query, all_tables):
    """ gets a list of all tables in the workspace and checks for the tables 
    used in the query
    
    Args:
        kql_query (str): KQL query
    
    Returns: list of query tables
    """
    query_tables = []
    for line in kql_query.splitlines():
        res = [tablename for tablename in all_tables if(
            re.search(f"{tablename}\\b", line) and "where" not in line and "|" not in line)
        ]
        if res:
            query_tables.append(res[0])

    query_tables = list(dict.fromkeys(query_tables))
    return query_tables

def create_word_doc(analytic_rules, all_tables):
    document = Document('NewDocTemplate.docx')

    for analytic_rule in analytic_rules:
        query_tables = []
        if analytic_rule.kind == "Scheduled":
            query_tables = get_query_tables(analytic_rule.query, all_tables)
        document.add_heading(analytic_rule.display_name, 1)
        table = document.add_table(rows=5, cols=2) # change number of table rows/columns here
        table.style = 'NewStyle' #change table style here
        row_cells = table.rows[0].cells
        row_cells[0].text = 'GUID'
        row_cells[0].width = Inches(1.06)
        row_cells[1].text = analytic_rule.name
        row_cells[1].width = Inches(5.42)
        row_cells = table.rows[1].cells
        row_cells[0].text = 'Description'
        row_cells[1].text = analytic_rule.description
        row_cells = table.rows[2].cells
        row_cells[0].text = 'Severity'
        if hasattr(analytic_rule, 'severity'):
            row_cells[1].text = analytic_rule.severity
        row_cells = table.rows[3].cells
        row_cells[0].text = 'Requirements'
        row_cells[1].text = ', '.join(query_tables)
        row_cells = table.rows[4].cells
        row_cells[0].text = 'Tactics'
        if hasattr(analytic_rule, 'tactics'):
            row_cells[1].text = ", ".join(analytic_rule.tactics)

        document.add_heading('Triage steps', 2)
        document.add_paragraph()
        document.add_heading('Change notes', 2)

        document.add_page_break()
    
    document.save('analytic_rules.docx')


def get_analytic_rules(args, securityinsights_client):
    """ get analytic rules from the Security Insights API
    
    Returns: list of analytic rules
    """
    rule_list = []

    try:
        analytic_rules = securityinsights_client.alert_rules.list(args.resource_group, args.workspace)
        print("[+] Getting analytic rules...")
        for analytic_rule in analytic_rules:
            if args.scheduled:
                if analytic_rule.kind == "Scheduled" in analytic_rule.display_name:
                    rule_list.append(analytic_rule)
            else:
                rule_list.append(analytic_rule)
        
        if args.enabled:
            rule_list = [rule for rule in rule_list if rule.enabled == True]

        rule_list.sort(key=lambda x: x.display_name)

        return rule_list
    except Exception as e:
        print(f"[-] Error getting analytic rules: {e}")
        sys.exit(1)


def get_table_names(args, log_analytics_client):
    """ get all table names and saved search names from the Log Analytics API

    Returns: list of table names
    """

    all_table_names = []
    try:
        tables = log_analytics_client.tables.list_by_workspace(
            args.resource_group, args.workspace)
        savedsearches=log_analytics_client.saved_searches.list_by_workspace(
            args.resource_group, args.workspace)
    except Exception as e:
        print(f"[-] Error getting tables: {e}")
        sys.exit(1)

    for table in tables:
        all_table_names.append(table.name)

    for savedsearch in savedsearches.value:
        if " " not in savedsearch.display_name:
            all_table_names.append(savedsearch.display_name)

    return all_table_names


def parse_args():
    parser = argparse.ArgumentParser(description="""
                                    Create a Microsoft Word document from 
                                    analytic rules in a Microsoft Sentinel 
                                    workspace""")
    parser.add_argument("-s", "--scheduled", action="store_true",
                        help="only include scheduled analytic rules")
    parser.add_argument("-e", "--enabled", action="store_true",
                        help="only include enabled analytic rules")
    parser.add_argument("-r", "--resource-group", required=True,
                        help="resource group name")
    parser.add_argument("-w", "--workspace", required=True,
                        help="workspace name")
    parser.add_argument("-i", "--subscription-id", required=True,
                        help="subscription id")
    args = parser.parse_args()

    try:
        securityinsights_client = SecurityInsights(
            credential=DefaultAzureCredential(), subscription_id=args.subscription_id)
        log_analytics_client = LogAnalyticsManagementClient(
            credential=DefaultAzureCredential(), subscription_id=args.subscription_id)
    except Exception as e:
        print(f"[-] Error creating Azure authentication client: {e}")
        sys.exit(1)

    rule_list = get_analytic_rules(args, securityinsights_client)
    all_tables = get_table_names(args, log_analytics_client)

    try:
        print("[+] Creating word document...")
        create_word_doc(rule_list, all_tables)
        print("[+] Word document created successfully")
    except Exception as e:
        print(f"[-] Error creating word document: {e}")
        sys.exit(1)


if __name__ == "__main__":
    args = parse_args()
